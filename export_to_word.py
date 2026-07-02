import os
import datetime
import re
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls
from docx.enum.section import WD_ORIENT

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def clean_special_characters(text):
    if not text:
        return ""
    # 1. Replace escaped pipes inside tables before splitting
    text = text.replace(r'\|', ' / ')
    
    # 1.3. Convert HTML line breaks to real newlines
    text = text.replace('<br>', '\n').replace('<br/>', '\n')
    
    # 1.5. Remove cent symbol (¢)
    text = text.replace(' ¢', '').replace('¢', '')
    
    # 2. Clean up math subscripts
    text = text.replace(r'$C_1$', 'C1')
    text = text.replace(r'$C_2$', 'C2')
    text = text.replace(r'$C_{DCA}$', 'CDCA')
    text = text.replace(r'$C_{swing}$', 'C_swing')
    text = text.replace(r'$', '')
    
    # 3. Strip emojis to avoid raw square blocks in Word
    emoji_pattern = re.compile(
        r'[\U00010000-\U0010ffff]|'  # Emojis and high planes
        r'[\u2600-\u27BF]|'          # Miscellaneous symbols / Dingbats
        r'[\u2300-\u23FF]|'          # Miscellaneous technical symbols
        r'[\u2B50-\u2B55]'           # Star, circle, etc.
    )
    text = emoji_pattern.sub('', text)
    
    # Clean up double spaces
    text = re.sub(r'  +', ' ', text)
    return text.strip()

def parse_markdown_to_doc(doc, md_text, PRIMARY_COLOR, SECONDARY_COLOR, HIGHLIGHT_COLOR):
    lines = md_text.split('\n')
    in_table = False
    table_rows = []
    
    i = 0
    while i < len(lines):
        line = clean_special_characters(lines[i].strip())
        
        # Check for empty line
        if not line:
            if in_table:
                build_table(doc, table_rows)
                in_table = False
                table_rows = []
            i += 1
            continue
            
        # Check for Heading 1 (e.g. # TITLE)
        if line.startswith('# '):
            if in_table:
                build_table(doc, table_rows)
                in_table = False
                table_rows = []
            h_text = line[2:]
            h = doc.add_heading(level=1)
            run = h.add_run(h_text)
            run.font.name = 'Arial'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = PRIMARY_COLOR
            h.alignment = WD_ALIGN_PARAGRAPH.CENTER
            h.paragraph_format.space_before = Pt(18)
            h.paragraph_format.space_after = Pt(10)
            h.paragraph_format.keep_with_next = True
            i += 1
            continue
            
        # Check for Heading 2 (e.g. ## SECTION)
        elif line.startswith('## '):
            if in_table:
                build_table(doc, table_rows)
                in_table = False
                table_rows = []
            h_text = line[3:]
            h = doc.add_heading(level=2)
            run = h.add_run(h_text)
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = PRIMARY_COLOR
            h.paragraph_format.space_before = Pt(14)
            h.paragraph_format.space_after = Pt(6)
            h.paragraph_format.keep_with_next = True
            i += 1
            continue
            
        # Check for Heading 3 (e.g. ### SUBSECTION)
        elif line.startswith('### '):
            if in_table:
                build_table(doc, table_rows)
                in_table = False
                table_rows = []
            h_text = line[4:]
            h = doc.add_heading(level=3)
            run = h.add_run(h_text)
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = SECONDARY_COLOR
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
            h.paragraph_format.keep_with_next = True
            i += 1
            continue
            
        # Check for Heading 4 (e.g. #### SUBSUBSECTION)
        elif line.startswith('#### '):
            if in_table:
                build_table(doc, table_rows)
                in_table = False
                table_rows = []
            h_text = line[5:]
            h = doc.add_heading(level=4)
            run = h.add_run(h_text)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            run.font.bold = True
            run.font.color.rgb = SECONDARY_COLOR
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(3)
            h.paragraph_format.keep_with_next = True
            i += 1
            continue
            
        # Check for Horizontal Rules
        elif line == '---' or line == '***':
            if in_table:
                build_table(doc, table_rows)
                in_table = False
                table_rows = []
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.space_before = Pt(12)
            p_format.space_after = Pt(12)
            # Add a subtle border or line
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), 'CCCCCC')
            pBdr.append(bottom)
            p._p.get_or_add_pPr().append(pBdr)
            i += 1
            continue
            
        # Check for Tables (line starting with '|')
        elif line.startswith('|'):
            in_table = True
            cols = [col.strip() for col in line.split('|')[1:-1]]
            # If this is the divider row (e.g. | :--- | :---: |), skip it
            if all(re.match(r'^[\s:-]+$', col) for col in cols):
                i += 1
                continue
            table_rows.append(cols)
            i += 1
            continue
            
        # Check for Bullet Lists (starts with '*' or '-')
        elif line.startswith('* ') or line.startswith('- '):
            if in_table:
                build_table(doc, table_rows)
                in_table = False
                table_rows = []
            
            p = doc.add_paragraph(style='List Bullet')
            p_format = p.paragraph_format
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(3)
            
            bullet_text = line[2:]
            parse_and_format_text(p, bullet_text, PRIMARY_COLOR, HIGHLIGHT_COLOR)
            i += 1
            continue
            
        # Normal text paragraph
        else:
            if in_table:
                build_table(doc, table_rows)
                in_table = False
                table_rows = []
                
            p = doc.add_paragraph()
            p_format = p.paragraph_format
            p_format.space_before = Pt(0)
            p_format.space_after = Pt(6)
            p_format.line_spacing = 1.15
            
            # Check if this is an italic metadata block
            if line.startswith('*') and line.endswith('*'):
                p_format.space_before = Pt(4)
                p_format.space_after = Pt(12)
                run = p.add_run(line.replace('*', ''))
                run.italic = True
                run.font.color.rgb = SECONDARY_COLOR
                run.font.size = Pt(9.5)
            else:
                parse_and_format_text(p, line, PRIMARY_COLOR, HIGHLIGHT_COLOR)
            
            i += 1

    # In case the table is at the very end
    if in_table:
        build_table(doc, table_rows)

def create_docx_report():
    print("--- STARTING WORD DOCUMENT EXPORT ---")
    
    # 1. Read CBOT_Reports_Log.md
    if not os.path.exists("CBOT_Reports_Log.md"):
        print("Error: CBOT_Reports_Log.md does not exist.")
        return
        
    with open("CBOT_Reports_Log.md", "r", encoding="utf-8") as f:
        md_content = f.read()

    # Define Colors
    PRIMARY_COLOR = RGBColor(0x1B, 0x36, 0x5D)  # Premium Dark Blue (#1B365D)
    SECONDARY_COLOR = RGBColor(0x5C, 0x76, 0x8D) # Slate Gray (#5C768D)
    HIGHLIGHT_COLOR = RGBColor(0xD9, 0x77, 0x06) # Warm Amber (#D97706)

    # ------------------ GENERATE FULL REPORT ------------------
    print("Generating Full Report...")
    doc = Document()
    
    # Page setup - Margins & Orientation - Landscape
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Inches(11.0)
        section.page_height = Inches(8.5)

    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    parse_markdown_to_doc(doc, md_content, PRIMARY_COLOR, SECONDARY_COLOR, HIGHLIGHT_COLOR)

    now_ict = datetime.datetime.now(datetime.timezone(datetime.timedelta(hours=7)))
    out_dir = "Daily Reports"
    os.makedirs(out_dir, exist_ok=True)
    filename = os.path.join(out_dir, now_ict.strftime("%d_%m_%Y-CBOT Reports.docx"))
    
    try:
        doc.save(filename)
        print(f"Successfully generated and saved Word report as: {filename}")
    except Exception as e:
        print(f"Error saving Daily Report: {e}")
    
    
    print("--- PROCESS COMPLETED ---")

def parse_and_format_text(paragraph, text, primary_color, highlight_color):
    """
    Parses Markdown bold '**text**', highlight '`code`', and emojis into the paragraph run objects.
    Also automatically highlights actual trade actions (Long/Short/Buy/Sell) in Green/Red and bold.
    """
    import re
    from docx.shared import RGBColor, Pt
    
    green_keywords = ["LỆNH LONG", "CANH MUA", "LỆNH BUY", "LONG", "BUY"]
    red_keywords = ["LỆNH SHORT", "CANH BÁN", "LỆNH SELL", "SHORT", "SELL"]
    
    keywords = sorted(green_keywords + red_keywords, key=len, reverse=True)
    escaped_keywords = [re.escape(kw) for kw in keywords]
    keyword_pattern = r'(' + '|'.join(escaped_keywords) + r')'
    
    pattern = r'(\*\*.*?\*\*|`.*?`)'
    tokens = re.split(pattern, text)
    
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            bold_val = token[2:-2]
            run = paragraph.add_run(bold_val)
            run.bold = True
            
            upper_bold = bold_val.upper()
            if any(kw in upper_bold for kw in [k.upper() for k in green_keywords]):
                run.font.color.rgb = RGBColor(0x15, 0x80, 0x3D) # Emerald Green
            elif any(kw in upper_bold for kw in [k.upper() for k in red_keywords]):
                run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C) # Red
        elif token.startswith('`') and token.endswith('`'):
            code_val = token[1:-1]
            run = paragraph.add_run(f" {code_val} ")
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = highlight_color
        else:
            sub_tokens = re.split(keyword_pattern, token, flags=re.IGNORECASE)
            for sub_token in sub_tokens:
                if not sub_token:
                    continue
                upper_sub = sub_token.upper()
                if upper_sub in [k.upper() for k in green_keywords]:
                    run = paragraph.add_run(sub_token)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0x15, 0x80, 0x3D) # Emerald Green
                elif upper_sub in [k.upper() for k in red_keywords]:
                    run = paragraph.add_run(sub_token)
                    run.bold = True
                    run.font.color.rgb = RGBColor(0xB9, 0x1C, 0x1C) # Red
                else:
                    paragraph.add_run(sub_token)

def build_table(doc, rows):
    """
    Creates a styled Word table from a list of rows with custom column widths (Landscape optimized).
    """
    if not rows:
        return
        
    num_cols = len(rows[0])
    num_rows = len(rows)
    
    table = doc.add_table(rows=num_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False # Disable autofit to allow manual cell sizing
    
    # 2. Determine column widths (Landscape width = 11 inches. Total table printable width ~ 9.5 inches)
    col_widths = []
    headers = [h.strip().lower() for h in rows[0]]
    
    if num_cols == 4 and any("vĩ mô" in h or "brent" in h for h in headers):
        # Macro Indicators Table
        col_widths = [Inches(1.8), Inches(1.2), Inches(1.2), Inches(5.3)]
    elif num_cols == 8 and any("lịch sử" in h or "ngày" in h for h in headers):
        # Report History Table
        col_widths = [Inches(1.1), Inches(0.6), Inches(1.0), Inches(1.4), Inches(0.8), Inches(1.2), Inches(2.4), Inches(1.0)]
    elif num_cols == 8 and any("chiến lược mùa vụ" in h or "vào lệnh" in h for h in headers):
        # Mua Vu Trade Matrix (8 cols)
        col_widths = [Inches(0.5), Inches(0.9), Inches(2.1), Inches(1.5), Inches(1.2), Inches(1.2), Inches(1.1), Inches(1.0)]
    elif num_cols == 7 and any("phân nhóm" in h or "tất toán" in h for h in headers):
        # Dual-Leg Trade Matrix (7 cols)
        col_widths = [Inches(0.9), Inches(1.8), Inches(1.5), Inches(1.1), Inches(1.3), Inches(1.9), Inches(1.0)]
    elif num_cols == 7 and any("chỉ tiêu" in h or "tiến độ" in h for h in headers):
        # Crop Progress Table (7 cols)
        col_widths = [Inches(0.7), Inches(0.9), Inches(1.0), Inches(1.0), Inches(1.0), Inches(1.5), Inches(3.4)]
    elif num_cols == 6 and any("nền tảng" in h or "chỉ tiêu" in h for h in headers):
        # Fundamentals Table
        col_widths = [Inches(1.6), Inches(0.8), Inches(1.0), Inches(0.9), Inches(1.2), Inches(4.0)]
    elif num_cols == 6 and any("chốt lời & tất toán" in h or "tất toán" in h for h in headers):
        # 6-column Mua Vu Trade Matrix with merged Entry/Exit timing
        col_widths = [Inches(1.2), Inches(2.3), Inches(1.8), Inches(1.2), Inches(2.0), Inches(1.0)]
    elif num_cols == 6 and any("thực chiến" in h or "chiến lược" in h for h in headers):
        # Action Summary Table
        col_widths = [Inches(0.6), Inches(1.5), Inches(1.8), Inches(1.4), Inches(1.4), Inches(2.8)]
    elif num_cols == 6 and any("matrix" in h or "managed money" in h for h in headers):
        # COT Matrix Table (6 cols)
        col_widths = [Inches(0.6), Inches(1.5), Inches(1.2), Inches(1.2), Inches(4.0), Inches(1.0)]
    elif num_cols == 6 and any("tồn kho" in h for h in headers):
        # WASDE Table (6 cols)
        col_widths = [Inches(0.6), Inches(1.1), Inches(1.6), Inches(1.6), Inches(1.2), Inches(3.4)]
    elif num_cols == 7 and any("xuất khẩu" in h or "bán hàng" in h for h in headers):
        # Export Table (7 cols)
        col_widths = [Inches(0.7), Inches(1.3), Inches(1.4), Inches(1.4), Inches(1.1), Inches(1.0), Inches(2.6)]
    elif num_cols == 5 and any("đối thủ" in h or "cung cầu" in h for h in headers):
        # Crop Production & Competitors Table (5 cols)
        col_widths = [Inches(0.7), Inches(1.8), Inches(2.0), Inches(1.2), Inches(3.8)]
    else:
        # Fallback distribution
        avg_width = Inches(9.5 / num_cols)
        col_widths = [avg_width] * num_cols
        
    # 3. Populate and format cells
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx]
        
        # Zebra striping and header styling
        is_header = (row_idx == 0)
        bg_color = "1B365D" if is_header else ("F8FAFC" if row_idx % 2 == 1 else "FFFFFF")
        
        for col_idx, text in enumerate(row_data):
            if col_idx >= num_cols:
                continue
                
            cell = row.cells[col_idx]
            cell.text = "" # Clear default text
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.05
            
            if is_header:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(text)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) # White text
                run.font.size = Pt(9.5)
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                parse_and_format_text(p, text, RGBColor(0x1B, 0x36, 0x5D), RGBColor(0xD9, 0x77, 0x06))
                
            # Formatting cell properties
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            
            # Apply width to this cell
            cell.width = col_widths[col_idx]
            
            # All cell borders (All Borders)
            tcPr = cell._tc.get_or_add_tcPr()
            tcBorders = parse_xml(
                f'<w:tcBorders {nsdecls("w")}>'
                f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
                f'</w:tcBorders>'
            )
            tcPr.append(tcBorders)

    # Spacing after table
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

if __name__ == '__main__':
    create_docx_report()
